import Search_ui
import CollectZoomeye
import ZoomeyeConfig
from PyQt5 import QtWidgets
from PyQt5 import *
import threading
import os,json
from docx import *
from urllib.parse import quote
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtWidgets import QMainWindow,QApplication,QHeaderView,QTableWidgetItem

class SearchInfomation(QtWidgets.QDialog):
    _SearchInfomation = pyqtSignal()
    _ZoomeyeThread = pyqtSignal()
    def __init__(self):
        QtWidgets.QWidget.__init__(self)
        self.ui = Search_ui.Ui_Form()
        self.ui.setupUi(self)
        self.fullchange = True
        self.i = 0

        self.ui.checkFull.stateChanged.connect(self.tableFull)
        self.ui.dataView.setContextMenuPolicy(Qt.CustomContextMenu)  # 添加右键功能

        self.stateFofa = False                            #radio改变状态调节，select为触发的函数吧
        self.stateZoomEye = False
        self.stateGoogle = False
        self.ui.Fofa.clicked.connect(self.select)
        self.ui.ZoomEye.clicked.connect(self.select)
        self.ui.Google.clicked.connect(self.select)
        self.ui.FofaConfig.setEnabled(False)
        self.ui.ZoomeyeConfig.setEnabled(False)
        self.ui.GoogleConfig.setEnabled(False)

        self.ZoomeyeConfig = ZoomeyeConfig.ZoomeyeConfig()

        self.jwtZoomeye = None
        self.keywordZoomeye = None
        self.CollectZoomeyeT = None
        #self.CollectZoomeyeT.filetrail.connect(self.writetrail)

        self.filename = None

    def tableFull(self):
        if self.fullchange == True:
            self.ui.dataView.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
            self.fullchange = False
        else:
            self.ui.dataView.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)
            self.ui.dataView.horizontalHeader().resizeSection(0, 80)
            self.ui.dataView.horizontalHeader().resizeSection(1, 80)
            self.ui.dataView.horizontalHeader().resizeSection(2, 80)
            self.ui.dataView.horizontalHeader().resizeSection(3, 80)
            self.ui.dataView.horizontalHeader().resizeSection(4, 80)
            self.fullchange = True

    def on_dataView_customContextMenuRequested(self, pos):  # 添加右键
        m = QMenu(self)
        m.addAction(self.ui.actionCheck_All)
        m.addAction(self.ui.actionCopy)
        m.popup(self.ui.dataView.mapToGlobal(pos))

    def select(self):
        if self.ui.Fofa.isChecked() and self.stateFofa == False:
            self.ui.FofaConfig.setEnabled(True)
            self.ui.ZoomeyeConfig.setEnabled(False)
            self.ui.GoogleConfig.setEnabled(False)
            print("选中fofa")
            self.stateFofa = True
        else:
            self.ui.FofaConfig.setEnabled(False)
            self.stateFofa = False
            self.ui.Fofa.setCheckable(False)
            self.ui.Fofa.setCheckable(True)
        if self.ui.ZoomEye.isChecked() and self.stateZoomEye == False:
            self.ui.FofaConfig.setEnabled(False)
            self.ui.ZoomeyeConfig.setEnabled(True)
            self.ui.GoogleConfig.setEnabled(False)
            print("选中ZoomEye")
            self.stateZoomEye = True
        else:
            self.ui.ZoomeyeConfig.setEnabled(False)
            self.stateZoomEye = False
            self.ui.ZoomEye.setCheckable(False)
            self.ui.ZoomEye.setCheckable(True)
        if self.ui.Google.isChecked() and self.stateGoogle == False:
            self.ui.FofaConfig.setEnabled(False)
            self.ui.ZoomeyeConfig.setEnabled(False)
            self.ui.GoogleConfig.setEnabled(True)
            print("选中Google")
            self.stateGoogle = True
        else:
            self.ui.GoogleConfig.setEnabled(False)
            self.stateGoogle = False
            self.ui.Google.setCheckable(False)
            self.ui.Google.setCheckable(True)

    @pyqtSlot(name="on_FofaConfig_clicked")
    def on_FofaConfig_clicked(self):
        print("hhh")

    @pyqtSlot(name="on_ZoomeyeConfig_clicked")
    def on_ZoomeyeConfig_clicked(self):
        ZoomeyeConfigres = self.ZoomeyeConfig.exec_()
        if ZoomeyeConfigres != QDialog.Accepted:
            return
        else:
            # if self.ZoomeyeConfig.jwtResult == None:
            #     print("未获取jwt")
            # else:
            #     self.jwtZoomeye = self.ZoomeyeConfig.jwtResult
            # if self.ZoomeyeConfig.keyword == '':
            #     print("未获取关键词")
            # else:
            #     self.keywordZoomeye = self.ZoomeyeConfig.keyword
            if self.ZoomeyeConfig.jwtResult == "":
                print("未获取jwt")
            else:
                self.jwtZoomeye = self.ZoomeyeConfig.jwtResult
                print("jwt"+self.jwtZoomeye)
            if self.ZoomeyeConfig.keyword == '':
                print("未获取关键词")
            else:
                self.keywordZoomeye = self.ZoomeyeConfig.keyword
                self.keywordZoomeye = self.keywordZoomeye.split()[0]
                print(self.keywordZoomeye)
                no = ["\\", "/", ":", "*", "?", "\"", "<", ">", "|"]
                self.filename = self.keywordZoomeye
                for i in range(0, len(no)):
                    if no[i] == "\"":
                        self.filename = self.filename.replace("%s" % no[i], "'")
                    self.filename = self.filename.replace("%s" % no[i], " ")
                try:
                    if os.path.exists("../SearchInformationFile/%s.doc" % self.filename):
                        os.remove("../SearchInformationFile/%s.doc" % self.filename)
                        doc = Document()
                        table = doc.add_table(rows=1, cols=6)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = ''
                        hdr_cells[1].text = 'timestamp'
                        hdr_cells[2].text = 'title'
                        hdr_cells[3].text = 'rdns'
                        hdr_cells[4].text = 'ip'
                        hdr_cells[5].text = 'service'
                        doc.save("../SearchInformationFile/%s.doc" % self.filename)
                    else:
                        doc = Document()
                        table = doc.add_table(rows=1, cols=6)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = ''
                        hdr_cells[1].text = 'timestamp'
                        hdr_cells[2].text = 'title'
                        hdr_cells[3].text = 'rdns'
                        hdr_cells[4].text = 'ip'
                        hdr_cells[5].text = 'service'
                        doc.save("../SearchInformationFile/%s.doc" % self.filename)
                    print("key"+self.keywordZoomeye)
                except Exception as e :
                    QMessageBox.information(self, '提示', '%s'%e)
            print("ZoomeyeConfig")

    @pyqtSlot(name="on_GoogleConfig_clicked")
    def on_GoogleConfig_clicked(self):
        print("hhhhh")



    @pyqtSlot(name='on_start_clicked')
    def on_start_clicked(self):
        # os.chdir("../common")
        # with open(r'proxyIP.py', 'w', encoding='utf-8') as f:
        #     f.write("import random\r\n")
        #     f.write("IPS = [\n")
        #     f.close()
        try:
            #self.start()
            if self.ui.Fofa.isChecked():
                print("Fofa")
            if self.ui.ZoomEye.isChecked():                           #初始化ZoomEye线程
                if self.jwtZoomeye == None:
                    QMessageBox.information(self, '提示', '请先登录ZoomEye获取认证~')
                else:
                    if self.keywordZoomeye == None:
                        QMessageBox.information(self, '提示', '请先输入关键词~')
                    else:
                        if self.CollectZoomeyeT == None:
                            print(self.jwtZoomeye, self.keywordZoomeye)
                            self.CollectZoomeyeT = CollectZoomeye.startCollectZoomeye(self.jwtZoomeye, quote(self.keywordZoomeye))
                            self.Zoomeyethread = QThread(self)
                            self.CollectZoomeyeT.moveToThread(self.Zoomeyethread)
                            self._ZoomeyeThread.connect(self.CollectZoomeyeT.run)
                            self.CollectZoomeyeT.collectZoomeye.connect(self.writeZoomeyecontent)
                            self.CollectZoomeyeT.overZoomeye.connect(self.overZoomeye)
                            self.startZoomeye()
                        else:
                            self.startZoomeye()
                print("ZoomEye")
            if self.ui.Google.isChecked():
                print("Google")
        except Exception as e:
            print("意外终止",e)


    def startZoomeye(self):
        if self.Zoomeyethread.isRunning():
            self.Zoomeyethread.quit()
            return
        self.Zoomeyethread.start()
        self._ZoomeyeThread.emit()

    def writeZoomeyecontent(self, timestamp, title, rdns, ip, service):
        rowPosotion = self.ui.dataView.rowCount()
        self.ui.dataView.insertRow(rowPosotion)
        self.ui.dataView.setItem(self.i, 0, QTableWidgetItem(timestamp))
        self.ui.dataView.setItem(self.i, 1, QTableWidgetItem(title))
        self.ui.dataView.setItem(self.i, 2, QTableWidgetItem(rdns))
        self.ui.dataView.setItem(self.i, 3, QTableWidgetItem(ip))
        self.ui.dataView.setItem(self.i, 4, QTableWidgetItem(service))
        doc = Document("../SearchInformationFile/%s.doc" % self.filename)
        table = doc.tables[0]
        table.add_row()
        hdr_cells = table.rows[self.i+1].cells
        hdr_cells[0].text = '%s'%(self.i+1)
        hdr_cells[1].text = '%s'%(timestamp)
        hdr_cells[2].text = '%s'%(title)
        hdr_cells[3].text = '%s'%(rdns)
        hdr_cells[4].text = '%s'%(ip)

        hdr_cells[5].text = '%s'%(service)
        doc.save("../SearchInformationFile/%s.doc" % self.filename)
        self.i = self.i + 1

    def overZoomeye(self):
        self.Zoomeyethread.exit()
        QMessageBox.about(self, '提示', 'ZoomEye收集完成,已保存在SearchInformationFile目录下~')

    @pyqtSlot(name='on_stop_clicked')
    def on_stop_clicked(self):
        if self.Zoomeyethread.isRunning():
            print("stopZoome")
            self.Zoomeyethread.terminate()
            QMessageBox.information(self, '提示', 'ZoomEye已经停止,已保存在SearchInformationFile目录下~')


    @pyqtSlot(name='on_actionCheck_All_triggered')
    def on_actionCheck_All_triggered(self):
        pass

    @pyqtSlot(name='on_actionCopy_triggered')
    def on_actionCopy_triggered(self):
        if self.i > 0:
            selected_ranges = self.ui.dataView.selectedRanges()[0]  # 只取第一个数据块,其他的如果需要要做遍历,简单功能就不写得那么复杂了
            text_str = ""  # 最后总的内容
            # 行（选中的行信息读取）
            for row in range(selected_ranges.topRow(), selected_ranges.bottomRow() + 1):
                row_str = ""
                # 列（选中的列信息读取）
                for col in range(selected_ranges.leftColumn(), selected_ranges.rightColumn() + 1):
                    item = self.ui.dataView.item(row, col)
                    row_str += item.text() + '\t'  # 制表符间隔数据
                text_str += row_str + '\n'  # 换行
            clipboard = qApp.clipboard()  # 获取剪贴板
            clipboard.setText(text_str)  # 内容写入剪贴板
    def keyPressEvent(self, event):
        """ Ctrl + C复制表格内容 """
        if self.i > 0:
            if event.modifiers() == Qt.ControlModifier and event.key() == Qt.Key_C:
                # 获取表格的选中行
                selected_ranges = self.ui.dataView.selectedRanges()[0]  # 只取第一个数据块,其他的如果需要要做遍历,简单功能就不写得那么复杂了
                text_str = ""  # 最后总的内容
                # 行（选中的行信息读取）
                for row in range(selected_ranges.topRow(), selected_ranges.bottomRow() + 1):
                    row_str = ""
                    # 列（选中的列信息读取）
                    for col in range(selected_ranges.leftColumn(), selected_ranges.rightColumn() + 1):
                        item = self.ui.dataView.item(row, col)
                        row_str += item.text() + '\t'  # 制表符间隔数据
                    text_str += row_str + '\n'  # 换行
                clipboard = qApp.clipboard()  # 获取剪贴板
                clipboard.setText(text_str)  # 内容写入剪贴板